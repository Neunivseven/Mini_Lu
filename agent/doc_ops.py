"""
办公文档编辑与排版：Word(.docx) / Excel(.xlsx) / PDF。

默认写到 data/docs_out/，并可对原文件旁生成 *_edited 副本，避免误覆盖。
复杂排版可通过 run_document_code 执行受限 Python。
"""
from __future__ import annotations

import json
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any

from agent.llm_client import app_dir

OUT_DIR_NAME = "docs_out"


def docs_out_dir() -> Path:
    p = app_dir() / "data" / OUT_DIR_NAME
    p.parent.mkdir(parents=True, exist_ok=True)
    p.mkdir(parents=True, exist_ok=True)
    return p


def resolve_doc_path(path: str, *, must_exist: bool = True) -> Path:
    p = Path(path).expanduser()
    if not p.is_absolute():
        # 相对路径：先项目根，再 docs_out
        cand = (app_dir() / p).resolve()
        if cand.exists() or not must_exist:
            p = cand
        else:
            p = (docs_out_dir() / p.name).resolve()
    else:
        p = p.resolve()
    if must_exist and not p.exists():
        raise FileNotFoundError(f"文件不存在: {p}")
    return p


def prepare_output(src: Path | None, *, suffix: str, filename: str = "") -> Path:
    """生成输出路径；若给 src 则默认旁路 *_edited。"""
    out = docs_out_dir()
    if filename:
        return out / filename
    if src is not None:
        stamp = datetime.now().strftime("%H%M%S")
        return out / f"{src.stem}_edited_{stamp}{suffix}"
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return out / f"document_{stamp}{suffix}"


def _backup_if_inplace(path: Path) -> Path | None:
    if not path.exists():
        return None
    bak = path.with_suffix(path.suffix + ".bak")
    shutil.copy2(path, bak)
    return bak


# ---------- 检视 ----------

def inspect_document(path: str) -> str:
    p = resolve_doc_path(path)
    suf = p.suffix.lower()
    lines = [f"文件: {p}", f"大小: {p.stat().st_size} 字节", f"类型: {suf}"]
    if suf == ".docx":
        lines.extend(_inspect_docx(p))
    elif suf in {".xlsx", ".xlsm"}:
        lines.extend(_inspect_xlsx(p))
    elif suf == ".pdf":
        lines.extend(_inspect_pdf(p))
    elif suf in {".txt", ".md", ".csv"}:
        text = p.read_text(encoding="utf-8", errors="replace")
        lines.append(f"行数: {text.count(chr(10)) + 1}，字符: {len(text)}")
    else:
        lines.append("（结构检视：该类型请用 read_document 看正文）")
    return "\n".join(lines)


def _inspect_docx(path: Path) -> list[str]:
    import docx

    doc = docx.Document(str(path))
    paras = [x.text for x in doc.paragraphs]
    non_empty = [t for t in paras if t.strip()]
    styles = {}
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        name = para.style.name if para.style else "Normal"
        styles[name] = styles.get(name, 0) + 1
    sec = doc.sections[0] if doc.sections else None
    margin = ""
    if sec:
        margin = (
            f"页边距(cm) 上{sec.top_margin.cm:.2f} 下{sec.bottom_margin.cm:.2f} "
            f"左{sec.left_margin.cm:.2f} 右{sec.right_margin.cm:.2f}"
        )
    lines = [
        f"段落数: {len(paras)}（非空 {len(non_empty)}）",
        f"表格数: {len(doc.tables)}",
        f"样式分布: {styles or '{}'}",
    ]
    if margin:
        lines.append(margin)
    preview = "\n".join(non_empty[:8])
    if preview:
        lines.append("--- 正文预览 ---")
        lines.append(preview[:1200])
    return lines


def _inspect_xlsx(path: Path) -> list[str]:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), data_only=True)
    lines = [f"工作表: {wb.sheetnames}"]
    for name in wb.sheetnames[:5]:
        ws = wb[name]
        lines.append(f"- {name}: {ws.dimensions}  max_row={ws.max_row} max_col={ws.max_column}")
        # 预览左上角
        rows = []
        for r in ws.iter_rows(min_row=1, max_row=min(6, ws.max_row or 1), max_col=min(6, ws.max_column or 1), values_only=True):
            rows.append(" | ".join("" if c is None else str(c) for c in r))
        if rows:
            lines.append("  预览:")
            lines.extend("  " + x for x in rows)
    return lines


def _inspect_pdf(path: Path) -> list[str]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    lines = [f"页数: {len(reader.pages)}"]
    if reader.metadata:
        meta = {k: str(v) for k, v in dict(reader.metadata).items() if v}
        if meta:
            lines.append(f"元数据: {meta}")
    for i, page in enumerate(reader.pages[:2]):
        t = (page.extract_text() or "").strip()
        if t:
            lines.append(f"--- 第{i+1}页预览 ---")
            lines.append(t[:800])
    return lines


# ---------- Word ----------

def edit_word(
    path: str,
    *,
    action: str,
    text: str = "",
    find: str = "",
    replace: str = "",
    style: str = "",
    font_name: str = "",
    font_size: float = 0,
    bold: bool | None = None,
    align: str = "",
    margin_cm: float = 0,
    output: str = "",
    inplace: bool = False,
) -> str:
    """
    action:
      replace_all | append | set_heading | format_all | set_margins | create
    """
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    action = (action or "").strip().lower()
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    if action == "create":
        out = resolve_doc_path(output, must_exist=False) if output else prepare_output(None, suffix=".docx")
        out.parent.mkdir(parents=True, exist_ok=True)
        doc = docx.Document()
        if style == "Title" or not style:
            doc.add_heading(text or "未命名文档", level=0)
        else:
            doc.add_paragraph(text or "")
        doc.save(str(out))
        return f"已创建 Word: {out}"

    src = resolve_doc_path(path)
    if src.suffix.lower() != ".docx":
        raise ValueError("edit_word 仅支持 .docx")

    if inplace:
        _backup_if_inplace(src)
        out = src
    else:
        out = resolve_doc_path(output, must_exist=False) if output else prepare_output(src, suffix=".docx")
        shutil.copy2(src, out)

    doc = docx.Document(str(out))

    if action == "replace_all":
        if not find:
            raise ValueError("replace_all 需要 find")
        n = 0
        for para in doc.paragraphs:
            if find in para.text:
                for run in para.runs:
                    if find in run.text:
                        run.text = run.text.replace(find, replace)
                        n += 1
                # 跨 run 简单兜底
                if find in para.text:
                    full = para.text.replace(find, replace)
                    if para.runs:
                        para.runs[0].text = full
                        for r in para.runs[1:]:
                            r.text = ""
                        n += 1
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if find in cell.text:
                        for para in cell.paragraphs:
                            if find in para.text:
                                for run in para.runs:
                                    run.text = run.text.replace(find, replace)
                                n += 1
        doc.save(str(out))
        return f"已替换约 {n} 处 → {out}"

    if action == "append":
        if style.lower().startswith("heading") or style in {"Title", "heading"}:
            level = 1
            m = re.search(r"(\d)", style)
            if m:
                level = int(m.group(1))
            doc.add_heading(text, level=min(level, 3))
        else:
            p = doc.add_paragraph(text)
            if style:
                try:
                    p.style = style
                except Exception:
                    pass
        doc.save(str(out))
        return f"已追加段落 → {out}"

    if action == "format_all":
        for para in doc.paragraphs:
            if align and align.lower() in align_map:
                para.alignment = align_map[align.lower()]
            for run in para.runs:
                if font_name:
                    run.font.name = font_name
                if font_size and font_size > 0:
                    run.font.size = Pt(font_size)
                if bold is not None:
                    run.bold = bold
        doc.save(str(out))
        return f"已应用全文排版 → {out}"

    if action == "set_margins":
        cm = margin_cm if margin_cm > 0 else 2.54
        for sec in doc.sections:
            sec.top_margin = Cm(cm)
            sec.bottom_margin = Cm(cm)
            sec.left_margin = Cm(cm)
            sec.right_margin = Cm(cm)
        doc.save(str(out))
        return f"已设置页边距 {cm}cm → {out}"

    raise ValueError(
        f"未知 action={action}，可用: create/replace_all/append/format_all/set_margins"
    )


# ---------- Excel ----------

def edit_excel(
    path: str = "",
    *,
    action: str,
    sheet: str = "",
    cells_json: str = "",
    text: str = "",
    font_name: str = "",
    font_size: float = 0,
    bold: bool | None = None,
    output: str = "",
    inplace: bool = False,
) -> str:
    """
    action:
      create | write_cells | format_range | append_row
    cells_json 示例: {"A1":"姓名","B1":"分数","A2":"Lee","B2":95}
    """
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Font

    action = (action or "").strip().lower()

    if action == "create":
        out = resolve_doc_path(output, must_exist=False) if output else prepare_output(None, suffix=".xlsx")
        out.parent.mkdir(parents=True, exist_ok=True)
        wb = Workbook()
        ws = wb.active
        ws.title = sheet or "Sheet1"
        if cells_json:
            cells = json.loads(cells_json)
            for addr, val in cells.items():
                ws[addr] = val
        elif text:
            ws["A1"] = text
        wb.save(str(out))
        return f"已创建 Excel: {out}"

    src = resolve_doc_path(path)
    if src.suffix.lower() not in {".xlsx", ".xlsm"}:
        raise ValueError("edit_excel 仅支持 .xlsx")

    if inplace:
        _backup_if_inplace(src)
        out = src
    else:
        out = resolve_doc_path(output, must_exist=False) if output else prepare_output(src, suffix=".xlsx")
        shutil.copy2(src, out)

    wb = load_workbook(str(out))
    ws = wb[sheet] if sheet and sheet in wb.sheetnames else wb.active

    if action == "write_cells":
        cells = json.loads(cells_json or "{}")
        if not cells:
            raise ValueError("write_cells 需要 cells_json")
        for addr, val in cells.items():
            ws[str(addr)] = val
        wb.save(str(out))
        return f"已写入 {len(cells)} 个单元格 → {out}"

    if action == "append_row":
        # text 为 JSON 数组，或逗号分隔
        if cells_json:
            row = json.loads(cells_json)
        else:
            row = [x.strip() for x in text.split(",")]
        if not isinstance(row, list):
            raise ValueError("append_row 需要数组")
        ws.append(row)
        wb.save(str(out))
        return f"已追加一行 → {out}"

    if action == "format_range":
        # cells_json 可为 {"range":"A1:C1"} 或地址列表
        spec = json.loads(cells_json or "{}")
        rng = spec.get("range") or spec.get("cells") or "A1"
        font = Font(
            name=font_name or None,
            size=font_size or None,
            bold=bold if bold is not None else None,
        )
        if ":" in str(rng) or isinstance(rng, str):
            for row in ws[rng]:
                for cell in row if hasattr(row, "__iter__") else [row]:
                    cell.font = font
        wb.save(str(out))
        return f"已设置格式 {rng} → {out}"

    raise ValueError(f"未知 action={action}，可用: create/write_cells/append_row/format_range")


# ---------- PDF ----------

def edit_pdf(
    path: str = "",
    *,
    action: str,
    text: str = "",
    pages: str = "",
    output: str = "",
    title: str = "",
) -> str:
    """
    action:
      create | extract_pages | merge（path 为 JSON 路径列表）
    PDF 精细排版建议用 run_document_code + reportlab。
    """
    action = (action or "").strip().lower()

    if action == "create":
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfgen import canvas

        out = resolve_doc_path(output, must_exist=False) if output else prepare_output(None, suffix=".pdf")
        out.parent.mkdir(parents=True, exist_ok=True)
        c = canvas.Canvas(str(out), pagesize=A4)
        width, height = A4
        # 尝试注册常见中文字体
        font_name = "Helvetica"
        for fp in (
            r"C:\Windows\Fonts\msyh.ttc",
            r"C:\Windows\Fonts\simsun.ttc",
            r"C:\Windows\Fonts\simhei.ttf",
        ):
            if Path(fp).exists():
                try:
                    pdfmetrics.registerFont(TTFont("CNFont", fp))
                    font_name = "CNFont"
                    break
                except Exception:
                    continue
        c.setFont(font_name, 14)
        y = height - 72
        if title:
            c.drawString(72, y, title)
            y -= 28
            c.setFont(font_name, 11)
        for line in (text or "").splitlines() or ["（空文档）"]:
            if y < 72:
                c.showPage()
                c.setFont(font_name, 11)
                y = height - 72
            c.drawString(72, y, line[:80])
            y -= 18
        c.save()
        return f"已创建 PDF: {out}"

    if action == "extract_pages":
        from pypdf import PdfReader, PdfWriter

        src = resolve_doc_path(path)
        reader = PdfReader(str(src))
        writer = PdfWriter()
        # pages: "1-3,5" 1-based
        indices = _parse_pages(pages, len(reader.pages))
        for i in indices:
            writer.add_page(reader.pages[i])
        out = resolve_doc_path(output, must_exist=False) if output else prepare_output(src, suffix=".pdf")
        with out.open("wb") as f:
            writer.write(f)
        return f"已导出页 {pages} → {out}"

    if action == "merge":
        from pypdf import PdfReader, PdfWriter

        paths = json.loads(path) if path.strip().startswith("[") else [path]
        writer = PdfWriter()
        for item in paths:
            rp = resolve_doc_path(str(item))
            reader = PdfReader(str(rp))
            for page in reader.pages:
                writer.add_page(page)
        out = resolve_doc_path(output, must_exist=False) if output else prepare_output(None, suffix=".pdf")
        with out.open("wb") as f:
            writer.write(f)
        return f"已合并 {len(paths)} 个 PDF → {out}"

    raise ValueError("未知 action，可用: create/extract_pages/merge；复杂排版请用 run_document_code")


def _parse_pages(spec: str, n_pages: int) -> list[int]:
    if not spec.strip():
        return list(range(n_pages))
    out: list[int] = []
    for part in spec.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            a, b = part.split("-", 1)
            start, end = int(a), int(b)
            out.extend(range(start - 1, end))
        else:
            out.append(int(part) - 1)
    return [i for i in out if 0 <= i < n_pages]
