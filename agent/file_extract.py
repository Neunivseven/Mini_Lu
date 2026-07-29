"""从本地办公文档提取纯文本，供 Agent 处理。"""
from __future__ import annotations

from pathlib import Path

SUPPORTED_SUFFIXES = {
    ".pdf",
    ".docx",
    ".xlsx",
    ".xlsm",
    ".txt",
    ".md",
    ".markdown",
    ".csv",
}
MAX_BYTES = 12 * 1024 * 1024  # 12MB
MAX_CHARS = 28000


def is_supported(path: str | Path) -> bool:
    return Path(path).suffix.lower() in SUPPORTED_SUFFIXES


def extract_text(path: str | Path, *, max_chars: int = MAX_CHARS) -> str:
    """提取文本；失败抛出带说明的异常。"""
    p = Path(path).expanduser().resolve()
    if not p.exists():
        raise FileNotFoundError(f"文件不存在: {p}")
    if not p.is_file():
        raise ValueError(f"不是文件: {p}")
    if p.stat().st_size > MAX_BYTES:
        raise ValueError(f"文件过大（>{MAX_BYTES // (1024*1024)}MB）: {p.name}")

    suf = p.suffix.lower()
    if suf == ".pdf":
        text = _from_pdf(p)
    elif suf == ".docx":
        text = _from_docx(p)
    elif suf in {".xlsx", ".xlsm"}:
        text = _from_xlsx(p)
    elif suf in {".txt", ".md", ".markdown", ".csv"}:
        text = _from_text(p)
    elif suf == ".doc":
        raise ValueError("暂不支持旧版 .doc，请另存为 .docx 或 PDF")
    else:
        raise ValueError(f"暂不支持的类型: {suf}（可用 pdf / docx / xlsx / txt）")

    text = (text or "").strip()
    if not text:
        raise ValueError(f"未能从文件提取到文字: {p.name}")
    max_chars = max(1000, min(int(max_chars), 80000))
    if len(text) > max_chars:
        return text[:max_chars] + f"\n\n…(已截断，原文约 {len(text)} 字)"
    return text


def build_agent_prompt(
    user_text: str,
    file_paths: list[str] | None = None,
    media_items: list[dict] | None = None,
) -> str:
    """把用户指令、文档附件、音/图识别结果拼成一轮 Agent 输入。"""
    user_text = (user_text or "").strip()
    file_paths = list(file_paths or [])
    media_items = list(media_items or [])
    parts: list[str] = []
    if user_text:
        parts.append(user_text)
    elif file_paths or media_items:
        parts.append("请结合下方附件与识别结果，给出简洁有用的回答。")
    else:
        parts.append("")

    if file_paths:
        parts.append("\n【附件】")
        for fp in file_paths:
            p = Path(fp)
            try:
                if p.suffix.lower() == ".pdf":
                    try:
                        from agent.doc_parsers import parse_document

                        result = parse_document(p, engine="auto")
                        body = result.get("text") or ""
                        eng = result.get("engine")
                        parts.append(
                            f"\n--- 文件: {p.name}（引擎 {eng}）---\n{body}\n--- 结束: {p.name} ---"
                        )
                        continue
                    except Exception:
                        pass
                body = extract_text(p)
                parts.append(f"\n--- 文件: {p.name} ---\n{body}\n--- 结束: {p.name} ---")
            except Exception as e:
                parts.append(f"\n--- 文件: {p.name} ---\n(读取失败: {e})\n---")

    if media_items:
        parts.append("\n【多模态识别结果】")
        for m in media_items:
            kind = str(m.get("kind") or "")
            name = str(m.get("name") or Path(str(m.get("path") or "")).name or "未命名")
            analysis = (m.get("analysis") or "").strip()
            label = {"audio": "语音识别", "image": "图像识别"}.get(kind, kind or "媒体")
            if analysis:
                parts.append(f"\n--- {label}: {name} ---\n{analysis}\n---")
            else:
                parts.append(f"\n--- {label}: {name} ---\n(无识别文本)\n---")

    return "\n".join(parts).strip() or "（空消息）"


def _from_text(path: Path) -> str:
    raw = path.read_bytes()
    for enc in ("utf-8", "utf-8-sig", "gbk", "gb18030", "latin-1"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _from_pdf(path: Path) -> str:
    # 优先 PyMuPDF；失败再 pypdf
    try:
        import fitz

        doc = fitz.open(str(path))
        try:
            chunks: list[str] = []
            for i in range(doc.page_count):
                t = (doc.load_page(i).get_text("text") or "").strip()
                if t:
                    chunks.append(f"[第{i + 1}页]\n{t}")
            return "\n\n".join(chunks)
        finally:
            doc.close()
    except ImportError:
        pass
    except Exception:
        pass
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError("缺少依赖 pymupdf 或 pypdf，请 pip install pymupdf") from e
    reader = PdfReader(str(path))
    chunks = []
    for i, page in enumerate(reader.pages):
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        if t.strip():
            chunks.append(f"[第{i + 1}页]\n{t.strip()}")
    return "\n\n".join(chunks)


def _from_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError as e:
        raise RuntimeError("缺少依赖 python-docx，请 pip install python-docx") from e
    document = docx.Document(str(path))
    paras = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
    # 表格文字
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                paras.append(" | ".join(cells))
    return "\n".join(paras)


def _from_xlsx(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as e:
        raise RuntimeError("缺少依赖 openpyxl，请 pip install openpyxl") from e
    wb = load_workbook(str(path), data_only=True)
    chunks: list[str] = []
    for name in wb.sheetnames:
        ws = wb[name]
        chunks.append(f"[工作表] {name}")
        for row in ws.iter_rows(max_row=min(200, ws.max_row or 1), values_only=True):
            vals = ["" if c is None else str(c) for c in row]
            if any(v.strip() for v in vals):
                chunks.append("\t".join(vals))
    return "\n".join(chunks)
